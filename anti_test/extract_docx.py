"""Extract docx content to txt for analysis."""
from docx import Document

doc = Document('d:/AI/paper/CellSam/anti_test/CellProfiler_deeplab.docx')

with open('d:/AI/paper/CellSam/anti_test/CellProfiler_deeplab_content.txt', 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        f.write(para.text + '\n')
    
    # Also extract tables
    for table in doc.tables:
        f.write('\n--- TABLE ---\n')
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            f.write(' | '.join(cells) + '\n')

print("Extracted to CellProfiler_deeplab_content.txt")

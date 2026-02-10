# [DEPRECATED] This script has been archived.
#
# Archived: 2026-02-10
# Reason: Superseded by unified inference core (Phase 0)
# Replacement entry points:
#   - Training:           src/train.py
#   - Oracle evaluation:  tools/standardized_inference.py
#   - E2E evaluation:     tools/evaluate_e2e.py
#   - Multi-model eval:   tools/comprehensive_eval.py
#   - Regression test:    tools/test_phase0_regression.py
#
import warnings as _warnings
_warnings.warn(
    "This script is deprecated. See header for replacement entry points.",
    DeprecationWarning, stacklevel=2
)
"""Extract docx content to txt for analysis."""
from docx import Document

doc = Document('d:/AI/paper/CellSam/anti_test/CellProfiler_deeplab.docx')

with open('d:/AI/paper/CellSam/anti_test/CellProfiler_deeplab_content.txt', 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        f.write(para.text + '\n')
    
    # Also extract tables
    for table in doc.tables:
        f.write('\n--- TABLE ---\n')
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            f.write(' | '.join(cells) + '\n')

print("Extracted to CellProfiler_deeplab_content.txt")
